"""
5W提问法 Word图卡生成器（通用版）

用法：
    python generate_5w.py "超市购物"           # 自动搜索图片 + 生成Word
    python generate_5w.py "公交车" --img /tmp/xxx.jpg  # 指定图片生成
    python generate_5w.py                      # 交互式选择已有场景
"""
import os, sys, json, re, subprocess, urllib.parse
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import html as html_mod

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SCENES_FILE = os.path.join(BASE_DIR, "5w_scenes.json")
IMG_DIR = os.path.join(BASE_DIR, "5w_images")
DESKTOP = "/mnt/c/Users/123/Desktop"

os.makedirs(IMG_DIR, exist_ok=True)

# ============ Word 样式工具 ============
def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, color=None, size=10, font="微软雅黑"):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font)
    rPr.insert(0, rFonts)
    return run

# ============ Bing图片搜索 ============
def bing_search(keywords, limit=8):
    """Bing图片搜索，返回真实图片URL列表（从murl字段提取）"""
    encoded = urllib.parse.quote(keywords)
    url = f"https://cn.bing.com/images/search?q={encoded}&first=1"
    raw = subprocess.run(
        ["curl", "-s", "--max-time", "10", "-A", "Mozilla/5.0", url],
        capture_output=True, text=True
    ).stdout
    # 解码HTML实体，从"murl"字段提取真实图片URL
    text = raw.replace('&quot;', '"').replace('&amp;', '&')
    imgs = re.findall(r'"murl":"(https://[^"]+\.(?:jpg|jpeg|png))"', text, re.IGNORECASE)
    seen, uniq = set(), []
    for img in imgs:
        if img not in seen:
            seen.add(img)
            uniq.append(img)
    return uniq[:limit]

def download_img(url, path):
    """curl下载图片到本地"""
    r = subprocess.run(
        ["curl", "-s", "--max-time", "15", "-A", "Mozilla/5.0", "-o", path, url],
        capture_output=True
    )
    return os.path.exists(path) and os.path.getsize(path) > 5000

# ============ 生成Word ============
def generate_5w_docx(topic, img_path, questions, output_path):
    """生成Word图卡（图片居中 + 5W问题表格）"""
    doc = Document()
    page = doc.sections[0]
    page.page_width = Cm(21)
    page.page_height = Cm(29.7)
    page.left_margin = page.right_margin = Cm(1.8)
    page.top_margin = page.bottom_margin = Cm(1.5)

    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    add_run(title, f"【{topic}】5W提问练习", bold=True, color="E8650A", size=16)

    # 图片
    if img_path and os.path.exists(img_path):
        pic_para = doc.add_paragraph()
        pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pic_para.add_run().add_picture(img_path, width=Inches(5.5))
        pic_para.paragraph_format.space_after = Pt(8)

    # 5W 问题表格
    W_COLORS = {
        "谁 Who":          ("谁",          "D6EAF8"),
        "什么 What":       ("什么",        "D5F5E3"),
        "哪里 Where":      ("哪里",        "FCF3CF"),
        "什么时候 When":    ("什么时候",    "FADBD8"),
        "为什么 Why":      ("为什么",      "E8DAEF"),
        "做什么 What doing":("做什么",      "FDEBD0"),
    }

    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False

    col_widths = [Cm(3.0), Cm(14.0)]
    for row in table.columns[0].cells:
        row.width = col_widths[0]
    for row in table.columns[1].cells:
        row.width = col_widths[1]

    for w_key, qs in questions.items():
        if not qs:
            continue
        label, bg = W_COLORS.get(w_key, (w_key.split()[0], "F0F0F0"))

        # 分类标签行（跨两列）
        lbl_row = table.add_row()
        lbl_cell = lbl_row.cells[0].merge(lbl_row.cells[1])
        lbl_cell.text = ""
        lp = lbl_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(2)
        lp.paragraph_format.left_indent = Pt(6)
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_run(lp, f"◆ {label}", bold=True, color="1A5F8A", size=11)
        set_cell_bg(lbl_cell, bg)

        # 问题行
        for q in qs:
            row = table.add_row()
            row.cells[0].text = ""
            row.cells[1].text = ""
            p0 = row.cells[0].paragraphs[0]
            p0.paragraph_format.space_before = Pt(1)
            p0.paragraph_format.space_after = Pt(1)
            p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p0, "□", size=12)
            p1 = row.cells[1].paragraphs[0]
            p1.paragraph_format.space_before = Pt(1)
            p1.paragraph_format.space_after = Pt(1)
            p1.paragraph_format.left_indent = Pt(4)
            add_run(p1, q, size=10)

    # 页脚
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(fp, "上海小螺号能力开发中心  ·  5W语言训练图卡  ·  tc-slp.netlify.app",
            color="888888", size=8)

    doc.save(output_path)
    return output_path

# ============ 交互式选择已有场景 ============
def list_scenes():
    if not os.path.exists(SCENES_FILE):
        return []
    with open(SCENES_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)
    return data.get("scenes", [])

def pick_scene():
    scenes = list_scenes()
    if not scenes:
        return None
    print("\n已有场景：")
    for i, s in enumerate(scenes, 1):
        print(f"  {i}. {s['name']}")
    print(f"  0. 退出\n")
    try:
        n = int(input("请选择场景编号: "))
        if n == 0:
            return None
        return scenes[n - 1]
    except:
        return None

# ============ 主函数 ============
def main():
    import argparse
    parser = argparse.ArgumentParser(description="5W提问法Word图卡生成器")
    parser.add_argument("topic", nargs="?", help="主题名称，如'公交车'")
    parser.add_argument("--img", help="指定本地图片路径")
    parser.add_argument("--url", help="指定图片URL")
    parser.add_argument("--questions", help="5W问题JSON字符串，格式:谁 Who:问题1|问题2;什么 What:...")
    args = parser.parse_args()

    topic = args.topic
    img_path = args.img

    # 无参数：交互式选择已有场景
    if not topic and not img_path:
        scene = pick_scene()
        if not scene:
            print("退出。")
            return
        topic = scene["name"]
        img_path = os.path.join(IMG_DIR, f"{scene['id']}.jpg")
        if not os.path.exists(img_path):
            img_path = None
        questions = scene.get("questions", {})
        out_file = os.path.join(DESKTOP, f"5W_{topic}.docx")
        result = generate_5w_docx(topic, img_path, questions, out_file)
        print(f"\n✅ 生成完成: {os.path.basename(result)}")
        return

    # 有主题参数：自动搜索图片
    if topic:
        print(f"\n🌈 5W提问法图卡生成")
        print(f"   主题: {topic}")

        # 下载图片
        if args.url:
            img_path = f"/tmp/auto_{topic}.jpg"
            ok = download_img(args.url, img_path)
            if ok:
                print(f"   📥 使用指定图片URL")
            else:
                print(f"   ❌ 图片下载失败")
                return
        elif not img_path:
            print(f"\n🔍 搜索图片中...")
            kws = [f"卡通 {topic} 小朋友", f"卡通 {topic} 儿童", f"儿童 {topic}"]
            urls = []
            for kw in kws:
                urls = bing_search(kw, limit=6)
                if urls:
                    print(f"   关键词「{kw}」找到 {len(urls)} 个候选")
                    break

            if not urls:
                print("   ❌ 未找到图片，尝试其他关键词")
                return

            img_path = f"/tmp/auto_{topic}.jpg"
            print(f"\n📥 尝试下载图片...")
            ok = False
            for url in urls[:4]:
                if download_img(url, img_path):
                    print(f"   ✅ 下载成功")
                    ok = True
                    break
            if not ok:
                print("   ❌ 所有候选图片均下载失败")
                return

        # 输出路径
        out_file = os.path.join(DESKTOP, f"5W_{topic}.docx")

        # 生成问题（基于图片描述）
        print(f"\n📋 图片已就绪: {img_path}")
        print(f"   请用 vision_analyze 分析图片：")
        print(f'   vision_analyze("{img_path}", "详细描述图中所有人物的外貌、穿着、动作、')
        print(f'   表情，以及场景里的物品和细节。用中文。")')
        print(f"\n⚠️  需要人工参与：AI描述图内容后，手动编写5W问题")
        print(f"   或使用 --questions 参数传入问题JSON")
        print(f"\n💡 完整流程：")
        print(f"   1. vision_analyze 分析图片")
        print(f"   2. 根据AI描述，在5w_scenes.json中添加问题")
        print(f"   3. 运行 python generate_5w.py 重新生成Word")
        print(f"\n📄 Word将保存到: {out_file}")

if __name__ == "__main__":
    main()
